import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

CURRENT_DATE = datetime.now().strftime("%B %d, %Y")

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_manual_docx():
    docx_path = "KoriePay_Confidential_System_Financial_Compliance_Manual_v1.0.docx"
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Tokens
    DARK_NAVY = RGBColor(8, 13, 26)
    EMERALD = RGBColor(16, 185, 129)
    TEAL = RGBColor(13, 148, 136)
    AMBER = RGBColor(217, 119, 6)
    ROSE = RGBColor(220, 38, 38)
    MUTED = RGBColor(71, 85, 105)

    # Header / Footer setup
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("KORIEPAY — STRICTLY CONFIDENTIAL // SYSTEM & FINANCIAL ARCHITECTURE MANUAL")
    hrun.font.size = Pt(8)
    hrun.font.color.rgb = MUTED

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    frun = fp.add_run("KoriePay Confidential — Unauthorized distribution, copying or disclosure is prohibited.   |   Strictly Confidential")
    frun.font.size = Pt(7.5)
    frun.font.color.rgb = MUTED

    # Cover Page
    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_badge = p_badge.add_run("KORIEPAY — STRICTLY CONFIDENTIAL // ACCESS LEVEL: AUTHORIZED PERSONNEL ONLY")
    r_badge.font.bold = True
    r_badge.font.size = Pt(9)
    r_badge.font.color.rgb = ROSE

    doc.add_paragraph()

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("KORIEPAY")
    r_title.font.bold = True
    r_title.font.size = Pt(28)
    r_title.font.color.rgb = DARK_NAVY

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Confidential System, Financial Architecture, Compliance & Implementation Manual")
    r_sub.font.bold = True
    r_sub.font.size = Pt(16)
    r_sub.font.color.rgb = TEAL

    p_desc = doc.add_paragraph()
    r_desc = p_desc.add_run("A Unified Reference & Engineering Blueprint for Tier-1 Core Banking, Double-Entry General Ledger, Central Liquidity Pools, Adashi / ROSCA Orchestration, and Bilateral Financial Rails across Nigeria (NGN) and Niger Republic (XOF).")
    r_desc.font.size = Pt(10)
    r_desc.font.color.rgb = MUTED

    # Metadata Table
    table_meta = doc.add_table(rows=10, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("DOCUMENT CONTROL IDENTIFIERS", "SPECIFICATION"),
        ("Document Classification", "STRICTLY CONFIDENTIAL // INTERNAL USE ONLY"),
        ("Access Level", "Tier-1 Leadership, Core Engineering, Treasury, Compliance, Regulated Partners"),
        ("Document Owner", "KoriePay Executive Management & Board of Directors"),
        ("Technical Owner", "KoriePay Technology, Core Banking & Cybersecurity Division"),
        ("Document Version / Status", "Version 1.0 — Approved Institutional Baseline"),
        ("Effective Publication Date", CURRENT_DATE),
        ("Operating Jurisdictions", "Nigeria (NGN / Providus Bank Node) & Niger Republic (XOF / Koris Bank Node)"),
        ("Mandatory Review Cycle", "Quarterly or upon material system, infrastructural, or regulatory changes"),
        ("Authoritative Truth Rule", "Double-Entry Core Ledger is Sole Financial Authority"),
    ]

    for idx, (k, v) in enumerate(meta_rows):
        row = table_meta.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.5)
        
        p_k = cell_k.paragraphs[0]
        p_v = cell_v.paragraphs[0]
        
        if idx == 0:
            set_cell_background(cell_k, "080D1A")
            set_cell_background(cell_v, "080D1A")
            rk = p_k.add_run(k)
            rk.font.bold = True
            rk.font.color.rgb = RGBColor(255, 255, 255)
            rv = p_v.add_run(v)
            rv.font.bold = True
            rv.font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_background(cell_k, "F8FAFC" if idx % 2 == 1 else "FFFFFF")
            set_cell_background(cell_v, "F8FAFC" if idx % 2 == 1 else "FFFFFF")
            rk = p_k.add_run(k)
            rk.font.bold = True
            rk.font.size = Pt(8.5)
            rv = p_v.add_run(v)
            rv.font.size = Pt(8.5)

    doc.add_page_break()

    # CHAPTER 1
    h1 = doc.add_paragraph()
    r = h1.add_run("1.0 MANDATORY REGULATORY & COMPLIANCE DISCLAIMER")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_NAVY

    p_disc = doc.add_paragraph()
    p_disc.add_run("This manual is an internal technology, operational, financial architecture, and compliance-readiness document prepared exclusively for authorized executives, engineers, auditors, and partner financial institutions of KoriePay. This document does not constitute formal legal advice, a direct banking licence, payment service bank licence, or standalone regulatory authorization. All financial operations in Nigeria operate in strict alignment with Central Bank of Nigeria (CBN) regulations and Providus Bank PLC clearing nodes. All operations in Niger Republic operate within the West African Economic and Monetary Union (WAEMU / UEMOA) framework supervised by the Banque Centrale des États de l'Afrique de l'Ouest (BCEAO) in partnership with Koris Bank SA.")

    # CHAPTER 2
    h2 = doc.add_paragraph()
    r = h2.add_run("2.0 CREDENTIALS & SECRETS MANAGEMENT ARCHITECTURE")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_NAVY

    p_sec = doc.add_paragraph()
    p_sec.add_run("CRITICAL ZERO-LEAKAGE INVARIANT: No plaintext passwords, production API secrets, database connection strings, JWT signing keys, or private bank credentials are published in this manual. All secrets are managed via AWS Secrets Manager and HashiCorp Vault with 90-day automated rotation and dual-custody access.")

    # CHAPTER 3
    h3 = doc.add_paragraph()
    r = h3.add_run("3.0 OPERATING JURISDICTIONS & BANKING NODES")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_NAVY

    table_jur = doc.add_table(rows=7, cols=3)
    jur_data = [
        ("OPERATING ATTRIBUTE", "NIGERIA (KP-NG)", "NIGER REPUBLIC (KP-NE)"),
        ("Legal Entity", "KoriePay Nigeria Limited (RC-1928392)", "KoriePay Niger SAS (RCCM-NI-NIA-2026-B-09)"),
        ("Currency", "Nigerian Naira (NGN, ₦)", "West African CFA Franc (XOF, CFA)"),
        ("Banking Node", "Providus Bank PLC (Clearing Vault)", "Koris Bank SA (Settlement Vault)"),
        ("Switch Integration", "NIBSS NIP / e-BillsPay", "GIM-UEMOA / BCEAO STAR-UEMOA"),
        ("Regulatory Authority", "Central Bank of Nigeria (CBN)", "BCEAO / WAEMU Regulatory Commission"),
        ("Identity Verification", "BVN & NIN Verification Gateway", "NIF & CNI National Identity Scheme"),
    ]
    for idx, (col1, col2, col3) in enumerate(jur_data):
        row = table_jur.rows[idx]
        for c_idx, text in enumerate([col1, col2, col3]):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            if idx == 0:
                set_cell_background(cell, "080D1A")
                run = p.add_run(text)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_background(cell, "F8FAFC" if idx % 2 == 1 else "FFFFFF")
                run = p.add_run(text)
                run.font.size = Pt(8.5)

    # CHAPTER 4
    h4 = doc.add_paragraph()
    r = h4.add_run("4.0 FINANCIAL TRUTH HIERARCHY & DOUBLE-ENTRY CORE LEDGER")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_NAVY

    p_led = doc.add_paragraph()
    p_led.add_run("The Double-Entry Core Ledger is the sole authoritative financial source of truth. Every transaction generates equal and opposite Debit and Credit journal entries: Total Assets = Total Liabilities + Total Equity. Product balances (wallets, float, Adashi) are downstream projections of the Core Ledger. Posted journal entries are strictly immutable.")

    # CHAPTER 5
    h5 = doc.add_paragraph()
    r = h5.add_run("5.0 ADASHI ROSCA SAVINGS ENGINE")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_NAVY

    p_ada = doc.add_paragraph()
    p_ada.add_run("Adashi (Ajo / Tontine) executes deterministic rotation allocation using HMAC-SHA256: HMAC_SHA256(Customer_UUID || Member_UUID, Group_UUID || Salt). Math.random() is strictly prohibited. Payout disbursements over ₦500,000 / 500,000 CFA enforce dual-authorization Maker-Checker approval.")

    # CHAPTER 6
    h6 = doc.add_paragraph()
    r = h6.add_run("6.0 CENTRAL LIQUIDITY POOL & NGN/XOF FIREWALL")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_NAVY

    p_liq = doc.add_paragraph()
    p_liq.add_run("NGN and XOF liquidity pools are strictly segregated. Cross-currency transfers require explicit FX Quote generation, approved rate locks, dual treasury signoff, and balanced double-entry ledger postings. Projected Available Liquidity = Opening Position + Confirmed Inflows + Expected Inflows - Confirmed Outflows - Expected Outflows - Reservations - Restrictions.")

    # Final Word
    doc.add_paragraph()
    p_final = doc.add_paragraph()
    p_final.add_run("FINAL ARCHITECTURAL PRINCIPLE: Identity establishes who the customer is; Risk & AML determine whether activity should proceed; Product Engines govern business rules; the Payment Switch orchestrates execution; the Core Ledger establishes absolute financial truth; Settlement executes external movement; Reconciliation verifies record agreement; Treasury manages liquidity; and the Immutable Audit Vault preserves institutional memory.")

    doc.save(docx_path)
    print(f"Master DOCX Manual generated successfully at: {docx_path}")

if __name__ == '__main__':
    build_manual_docx()
